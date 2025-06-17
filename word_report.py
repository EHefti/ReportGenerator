import os
import re
from docx import Document
from docx.shared import Inches


def create_word_report_tryexcept(title, stats_table_png_path, overview_paths, scatter_plot_paths, histogram_paths, success_paths, corr_table_paths, output_filename="report.docx", remove_files=False):
    """
    Generates a Word report from PNG image paths, handling separate paths for different plot types.

    Args:
        stats_table_png_path (str): Path to the PNG file of the statistics table.
        overview_paths (list): List of paths to the PNG files for overview plots (e.g., line, violin, bar).
        scatter_plot_paths (list): List of paths to the PNG files of the scatter plots.
        histogram_paths (list): List of paths to the PNG files of the histograms.
        success_paths (list): List of paths to the PNG files of the success chance plots.
        corr_table_paths (list): List of paths to the PNG files of the correlation tables.
        output_filename (str, optional): The name of the output Word file. Defaults to "report.docx".
        remove_files (bool, optional): Whether to remove the PNG files after creating the report. Defaults to False.
    """
    document = Document()

    # Titelseite
    document.add_heading(title, level=1)

    # 1. Stats Table
    document.add_heading("Statistische Zusammenfassung", level=2)
    document.add_paragraph("Auswertung der Sim. Prüfungsdaten von der Saison 24/25. Zuerst wird ein Überblick über die Serien gegeben, " \
                            "dann werden die Serien der SimPrüfung einzeln dargestellt. Eine Prüfung gilt als bestanden, wenn " \
                            "in der ZAP-Prüfung mindestens eine 4.5 geschrieben wurde.")

    try:
        if os.path.exists(stats_table_png_path):
            if re.search(r"glc", stats_table_png_path):
                document.add_picture(stats_table_png_path, width=Inches(4.4)) # , height=Inches(2.3)
            else:
                document.add_picture(stats_table_png_path, width=Inches(6)) # , height=Inches(1.4)
            if remove_files:
                os.remove(stats_table_png_path)
    except Exception as e:
        print(f"Fehler beim Hinzufügen der Statistik-Tabelle {stats_table_png_path}: {e}")

    document.add_paragraph("Legende:")
    document.add_paragraph(f"PNx (Alle) = Prüfungsnote SimPr (Alle SuS Serie x)", style='List Bullet')
    document.add_paragraph(f"PNx (mit EN) = Prüfungsnote SimPr (SuS mit Erfolgsnote Serie x)", style='List Bullet')
    document.add_paragraph(f"ENx = Erfolgsnote ZAP (SuS mit Erfolgsnote Serie x)", style='List Bullet')
    document.add_paragraph(f"PNall (Alle) = Prüfungsnote aller SimPr (Alle SuS aller Serien)", style='List Bullet')
    document.add_paragraph(f"PNall (mit EN) = Prüfungsnote aller SimPr (SuS mit Erfolgsnote in einer der Serien)", style='List Bullet')
    document.add_paragraph(f"ENall = Erfolgsnote ZAP (SuS mit Note in einer der Serien)", style='List Bullet')
    document.add_paragraph(f"EN = Alle Erfolgsnoten", style='List Bullet')

    document.add_page_break()

    # Übersicht Serien
    document.add_heading("Übersicht Serien", level=2)

    # Overview Plot 1 (Line Plot / First overview)
    if overview_paths and len(overview_paths) > 0:
        try:
            if os.path.exists(overview_paths[0]):
                document.add_picture(overview_paths[0], width=Inches(6)) # , height=Inches(4)
                document.add_paragraph("Die Boxplots zeigen den Mittelwert und die Verteilung der aller Prüfungsnoten der SimPrüfung, nur die Noten der Schüler mit einer guten Erfolgsnote," \
                                       "und die Prüfungsnoten der ZAP Prüfung. ")
                if remove_files:
                    os.remove(overview_paths[0])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 1 {overview_paths[0]}: {e}")


    # Overview Plot 2 (Violin Plot)
    if overview_paths and len(overview_paths) > 1:
        try:
            if os.path.exists(overview_paths[1]):
                document.add_picture(overview_paths[1], width=Inches(6)) # , height=Inches(4)
                document.add_paragraph("Ein Violin-Plot zeigt die Verteilung von Daten als Kombination eines Box-Plots und eines Dichte-Diagramm. " \
                                       "Der weisse Linie in der Mitte repräsentiert den Median der Daten (der mittlere Wert), während der dickere schwarze Balken die mittleren 50% der Daten anzeigt. " \
                                       "Der dünne schwarze Balken zeigt die restlichen Daten an (ohne statistische Ausreisser). " \
                                       "Die 'Geigenform' selbst zeigt die Dichte der Daten: Je breiter der Violin-Plot an einer Stelle ist, desto mehr Datenpunkte befinden sich in diesem Wertebereich. ")
                if remove_files:
                    os.remove(overview_paths[1])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 2 (Violin-Plot) {overview_paths[1]}: {e}")


    # Overview Plot 3 (Bar Plot)
    if overview_paths and len(overview_paths) > 2:
        try:
            if os.path.exists(overview_paths[2]):
                document.add_picture(overview_paths[2], width=Inches(6)) # , height=Inches(4)
                document.add_paragraph("Das Histogramm representiert die Bestehensquote der einzelnen Serien, verglichen mit der ZAP-Prüfung.")
                if remove_files:
                    os.remove(overview_paths[2])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 3 (Balken-Plot) {overview_paths[2]}: {e}")


    # Overview Plot 4 (Success Chance all Series)
    if overview_paths and len(overview_paths) > 3:
        try:
            if os.path.exists(overview_paths[3]):
                document.add_picture(overview_paths[3], width=Inches(6)) # , width=Inches(6), height=Inches(4))
                document.add_paragraph("Der Line-Plot zeigt die Erfolgsquote mit dem Durchschnitt aller Serien, verglichen mit der ZAP-Prüfung.")
                if remove_files:
                    os.remove(overview_paths[3])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 3 (Balken-Plot) {overview_paths[2]}: {e}")

    document.add_page_break()

    # 2. Serien einzeln
    document.add_heading("Auswertung der Serien", level=2)
    
    # Iterate through individual series plots
    # It's safer to iterate based on the longest list or a known number of series,
    # then check existence for each specific plot.
    num_series = max(len(scatter_plot_paths), len(histogram_paths), len(success_paths), len(corr_table_paths))

    for i in range(num_series):
        # Add sub-heading for series if any plot exists for it
        # This prevents an empty "Serie X" heading if all plots for it are missing
        series_plots_exist = False
        if (i < len(scatter_plot_paths) and os.path.exists(scatter_plot_paths[i])) or \
           (i < len(histogram_paths) and os.path.exists(histogram_paths[i])) or \
           (i < len(success_paths) and os.path.exists(success_paths[i])) or \
           (i < len(corr_table_paths) and os.path.exists(corr_table_paths[i])):
            document.add_heading(f"Serie {i + 1}", level=3)
            series_plots_exist = True
        elif i < len(scatter_plot_paths) or i < len(histogram_paths) or i < len(success_paths) or i < len(corr_table_paths):
            # print(f"Hinweis: Keine Plots für Serie {i+1} gefunden, Sektion wird übersprungen.")
            continue # Skip to next series if no plots exist for this one

        # Scatterplot
        if i < len(scatter_plot_paths):
            try:
                if os.path.exists(scatter_plot_paths[i]):
                    document.add_heading(f"Scatterplot: Serie {i + 1}", level=4)
                    document.add_paragraph(f"Folgender Scatterplot vergleicht die Prüfungsnote der Serie {i+1} und der ZAP-Prüfung."
                                            "Alle Schüler, die die Serie und die ZAP-Prüfung abgeschlossen haben, werden in diesem Plot dargestellt")
                    document.add_picture(scatter_plot_paths[i], width=Inches(6)) #, height=Inches(4.2)
                    if remove_files:
                        os.remove(scatter_plot_paths[i])
                    document.add_page_break() # Add page break after each scatter plot
            except Exception as e:
                print(f"Fehler beim Hinzufügen des Scatterplots für Serie {i+1} ({scatter_plot_paths[i]}): {e}")

        # Histograms
        if i < len(histogram_paths):
            try:
                if os.path.exists(histogram_paths[i]):
                    document.add_heading(f"Histogramme: Serie {i + 1}", level=4)
                    document.add_paragraph(f"Folgende Histogramme zeigen die Verteilung der Prüfungsnote der Serie {i+1}."
                                            "Dabei sieht man wie sich die SuS mit guten guten Erfolgsnoten von dem allgemeinen Notenbild abheben.")
                    document.add_picture(histogram_paths[i], width=Inches(6)) # , height=Inches(3.6)
                    if remove_files:
                        os.remove(histogram_paths[i])
            except Exception as e:
                print(f"Fehler beim Hinzufügen des Histogramms für Serie {i+1} ({histogram_paths[i]}): {e}")

        # Success Chance Plot
        if i < len(success_paths):
            try:
                if os.path.exists(success_paths[i]):
                    document.add_heading(f"Erfolgsquote: Serie {i + 1}", level=4)
                    document.add_paragraph(f"Folgender Plot zeigt die Erfolgsquote in der ZAP Prüfung basierend auf der SimPrüfung der Serie {i+1}.")
                    document.add_picture(success_paths[i], width=Inches(6)) # , height=Inches(3.6)
                    if remove_files:
                        os.remove(success_paths[i])
            except Exception as e:
                print(f"Fehler beim Hinzufügen des Erfolgsquote-Plots für Serie {i+1} ({success_paths[i]}): {e}")

        # Correlation Table
        if i < len(corr_table_paths):
            try:
                if os.path.exists(corr_table_paths[i]):
                    document.add_heading(f"Korrelationstabelle: Serie {i + 1}", level=4)
                    document.add_paragraph(f"Folgende Tabelle zeigt die Korrelationen zwischen den verschiedenen Aufgaben und der Prüfungsnote der Serie {i+1}.")
                    if re.search(r"glc", corr_table_paths[i]):
                        document.add_picture(corr_table_paths[i], width=Inches(4)) # , height=Inches(2.3)
                    else:
                        document.add_picture(corr_table_paths[i], width=Inches(6)) # , height=Inches(1.7)
                    if remove_files:
                        os.remove(corr_table_paths[i])
            except Exception as e:
                print(f"Fehler beim Hinzufügen der Korrelationstabelle für Serie {i+1} ({corr_table_paths[i]}): {e}")

        if series_plots_exist: # Only add page break if some plot was actually added for this series
             document.add_page_break()

    # Save the Word report
    try:
        document.save(output_filename)
        print(f"Word report '{output_filename}' created successfully!")
    except Exception as e:
        print(f"Fehler beim Speichern des Word-Reports '{output_filename}': {e}")



def create_word_report_tryexcept_glc(title, stats_table_png_path, overview_paths, scatter_plot_paths, histogram_paths, success_paths, corr_table_paths, output_filename="report.docx", remove_files=False):
    """
    Generates a Word report from PNG image paths, handling separate paths for different plot types.

    Args:
        stats_table_png_path (str): Path to the PNG file of the statistics table.
        overview_paths (list): List of paths to the PNG files for overview plots (e.g., line, violin, bar).
        scatter_plot_paths (list): List of paths to the PNG files of the scatter plots.
        histogram_paths (list): List of paths to the PNG files of the histograms.
        success_paths (list): List of paths to the PNG files of the success chance plots.
        corr_table_paths (list): List of paths to the PNG files of the correlation tables.
        output_filename (str, optional): The name of the output Word file. Defaults to "report.docx".
        remove_files (bool, optional): Whether to remove the PNG files after creating the report. Defaults to False.
    """
    document = Document()

    # Titelseite
    document.add_heading(title, level=1)

    # 1. Stats Table
    document.add_heading("Statistische Zusammenfassung", level=2)
    document.add_paragraph("Auswertung der GLC Prüfungsdaten von der Saison 24/25. Eine Prüfung gilt als bestanden, wenn " \
                            "in der ZAP-Prüfung mindestens eine 4.5 geschrieben wurde. Um die GLC Prüfung mit der ZAP Prüfung zu vergleichen, " \
                            "wurde die GLC Prüfung mit 6 multipliziert, so haben wir die gleiche Skala für beide Prüfungen.")

    try:
        if os.path.exists(stats_table_png_path):
            if re.search(r"glc", stats_table_png_path):
                document.add_picture(stats_table_png_path, width=Inches(6)) # , width=Inches(4.4), height=Inches(2.3)
            else:
                document.add_picture(stats_table_png_path, width=Inches(6)) # , width=Inches(6), height=Inches(1.4)
            if remove_files:
                os.remove(stats_table_png_path)
    except Exception as e:
        print(f"Fehler beim Hinzufügen der Statistik-Tabelle {stats_table_png_path}: {e}")

    document.add_page_break()

    # Übersicht Serien
    document.add_heading("Übersicht", level=2)
    document.add_paragraph("Hier wird ein Überblick über die GLC Prüfung gegeben. Um die Noten vergleichen zu können, wurde die ZAP Note durch 6 geteilt.")

    # Overview Plot 1 (Line Plot / First overview)
    if overview_paths and len(overview_paths) > 0:
        try:
            if os.path.exists(overview_paths[0]):
                document.add_picture(overview_paths[0], width=Inches(6)) # , width=Inches(6), height=Inches(4))
                document.add_paragraph("Der Plot zeigt den Mittelwert und die Varianz, der Prüfungsnoten des GLC.")
                if remove_files:
                    os.remove(overview_paths[0])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 1 {overview_paths[0]}: {e}")


    # Overview Plot 2 (Violin Plot)
    if overview_paths and len(overview_paths) > 1:
        try:
            if os.path.exists(overview_paths[1]):
                document.add_picture(overview_paths[1], width=Inches(6)) # , width=Inches(6), height=Inches(4))
                document.add_paragraph("Ein Violin-Plot zeigt die Verteilung von Daten als Kombination eines Box-Plots und eines Dichte-Diagramm. " \
                                       "Der weisse Linie in der Mitte repräsentiert den Median der Daten (der mittlere Wert), während der dickere schwarze Balken die mittleren 50% der Daten anzeigt. " \
                                       "Der dünne schwarze Balken zeigt die restlichen Daten an (ohne statistische Ausreisser). " \
                                       "Die 'Geigenform' selbst zeigt die Dichte der Daten: Je breiter der Violin-Plot an einer Stelle ist, desto mehr Datenpunkte befinden sich in diesem Wertebereich. "\
                                       "Die EN wurde hier durch 6 geteilt um eine Vergleich im selben Plot zu ermöglichen.")
                if remove_files:
                    os.remove(overview_paths[1])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 2 (Violin-Plot) {overview_paths[1]}: {e}")

    
    
    # Overview Plot 3 (Bar Plot)
    if overview_paths and len(overview_paths) > 2:
        try:
            if os.path.exists(overview_paths[2]):
                """
                document.add_picture(overview_paths[2], width=Inches(6)) # , width=Inches(6), height=Inches(4))
                document.add_paragraph("Das Histogramm representiert die Bestehensquote der einzelnen Serien, verglichen mit der ZAP-Prüfung.")
                """
                if remove_files:
                    os.remove(overview_paths[2])
        except Exception as e:
            print(f"Fehler beim Hinzufügen von Überblick-Plot 3 (Balken-Plot) {overview_paths[2]}: {e}")

    
    document.add_page_break()
    
    
    # Iterate through individual series plots
    # It's safer to iterate based on the longest list or a known number of series,
    # then check existence for each specific plot.
    num_series = max(len(scatter_plot_paths), len(histogram_paths), len(success_paths), len(corr_table_paths))

    for i in range(num_series):
        # Add sub-heading for series if any plot exists for it
        # This prevents an empty "Serie X" heading if all plots for it are missing
        series_plots_exist = False
        if (i < len(scatter_plot_paths) and os.path.exists(scatter_plot_paths[i])) or \
           (i < len(histogram_paths) and os.path.exists(histogram_paths[i])) or \
           (i < len(success_paths) and os.path.exists(success_paths[i])) or \
           (i < len(corr_table_paths) and os.path.exists(corr_table_paths[i])):
            document.add_heading(f"GLC", level=2)
            series_plots_exist = True
        elif i < len(scatter_plot_paths) or i < len(histogram_paths) or i < len(success_paths) or i < len(corr_table_paths):
            # print(f"Hinweis: Keine Plots für Serie {i+1} gefunden, Sektion wird übersprungen.")
            continue # Skip to next series if no plots exist for this one

        # Scatterplot
        if i < len(scatter_plot_paths):
            try:
                if os.path.exists(scatter_plot_paths[i]):
                    document.add_heading(f"Scatterplot:", level=4)
                    document.add_paragraph(f"Folgender Scatterplot vergleicht die Prüfungsnote des GLC und der ZAP-Prüfung."
                                            "Alle Schüler, die die Serie und die ZAP-Prüfung abgeschlossen haben, werden in diesem Plot dargestellt")
                    document.add_picture(scatter_plot_paths[i], width=Inches(6)) # , width=Inches(6), height=Inches(4.2))
                    if remove_files:
                        os.remove(scatter_plot_paths[i])
                    document.add_page_break() # Add page break after each scatter plot
            except Exception as e:
                print(f"Fehler beim Hinzufügen des Scatterplots für Serie {i+1} ({scatter_plot_paths[i]}): {e}")

        # Histograms
        if i < len(histogram_paths):
            try:
                if os.path.exists(histogram_paths[i]):
                    document.add_heading(f"Histogramme:", level=4)
                    document.add_paragraph(f"Folgende Histogramme zeigen die Verteilung der Prüfungsnote des GLC."
                                            "Dabei sieht man wie sich die SuS mit guten guten Erfolgsnoten von dem allgemeinen Notenbild abheben.")
                    document.add_picture(histogram_paths[i], width=Inches(6)) #, width=Inches(6), height=Inches(3.6)
                    if remove_files:
                        os.remove(histogram_paths[i])
            except Exception as e:
                print(f"Fehler beim Hinzufügen des Histogramms ({histogram_paths[i]}): {e}")

        # Success Chance Plot
        if i < len(success_paths):
            try:
                if os.path.exists(success_paths[i]):
                    document.add_heading(f"Erfolgsquote:", level=4)
                    document.add_paragraph(f"Folgender Plot zeigt die Erfolgsquote in der ZAP Prüfung basierend auf dem GLC.")
                    document.add_picture(success_paths[i], width=Inches(6)) # , width=Inches(6), height=Inches(3.6))
                    if remove_files:
                        os.remove(success_paths[i])
            except Exception as e:
                print(f"Fehler beim Hinzufügen des Erfolgsquote-Plots ({success_paths[i]}): {e}")

        # Correlation Table
        if i < len(corr_table_paths):
            try:
                if os.path.exists(corr_table_paths[i]):
                    document.add_heading(f"Korrelationstabelle:", level=4)
                    document.add_paragraph(f"Folgende Tabelle zeigt die Korrelationen zwischen den verschiedenen Aufgaben und der Prüfungsnote des GLC.")
                    if re.search(r"glc", corr_table_paths[i]):
                        document.add_picture(corr_table_paths[i], width=Inches(6)) # , width=Inches(4), height=Inches(2.3))
                    else:
                        document.add_picture(corr_table_paths[i], width=Inches(6)) # , width=Inches(6), height=Inches(1.7)
                    if remove_files:
                        os.remove(corr_table_paths[i])
            except Exception as e:
                print(f"Fehler beim Hinzufügen der Korrelationstabelle für Serie {i+1} ({corr_table_paths[i]}): {e}")

        if series_plots_exist: # Only add page break if some plot was actually added for this series
             document.add_page_break()

    # Save the Word report
    try:
        document.save(output_filename)
        print(f"Word report '{output_filename}' created successfully!")
    except Exception as e:
        print(f"Fehler beim Speichern des Word-Reports '{output_filename}': {e}")